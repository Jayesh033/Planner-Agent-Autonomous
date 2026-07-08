"""
Document builder: converts the executed plan into a polished Microsoft Word (.docx) file.

Formatting conventions:
  - Cover page with title, document type, date, and watermark-style subtitle.
  - Executive Summary (objective + assumptions) on page 2.
  - One section per task output with a numbered heading.
  - Consistent styles: Heading 1, Heading 2, Normal, Quote.
  - Page numbers in footer.
  - Reflection quality badge in the document properties (custom property).
"""
import logging
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm

from agent.models import ExecutionPlan, ReflectionResult

logger = logging.getLogger(__name__)

BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)    # near-black navy
BRAND_ACCENT = RGBColor(0x16, 0x21, 0x3E)   # deep blue
BRAND_LIGHT = RGBColor(0xE9, 0xEA, 0xEC)    # light grey
TEXT_SECONDARY = RGBColor(0x55, 0x5F, 0x6E) # muted text


def build_document(plan: ExecutionPlan, reflection: ReflectionResult, output_dir: Path) -> Path:
    """Build and save the Word document. Returns the path to the saved file."""
    doc = Document()
    _configure_page(doc)
    _add_cover_page(doc, plan, reflection)
    _add_executive_summary(doc, plan)
    _add_task_sections(doc, plan)
    _add_reflection_appendix(doc, reflection)
    _add_footer(doc, plan)

    filename = _safe_filename(plan.document_title) + ".docx"
    output_path = output_dir / filename
    doc.save(str(output_path))
    logger.info("Document saved: %s", output_path)
    return output_path


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)



def _add_cover_page(doc: Document, plan: ExecutionPlan, reflection: ReflectionResult) -> None:
    _add_spacer(doc, lines=4)

    # Document type label
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_type.add_run(plan.document_type.upper())
    run.font.size = Pt(12)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.bold = True

    _add_spacer(doc, lines=1)

    # Main title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(plan.document_title)
    run.font.size = Pt(26)
    run.font.color.rgb = BRAND_DARK
    run.font.bold = True

    _add_spacer(doc, lines=2)

    # Divider line
    _add_horizontal_rule(doc)

    _add_spacer(doc, lines=1)

    # Objective
    p_obj = doc.add_paragraph()
    p_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_obj.add_run(plan.objective)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.italic = True

    _add_spacer(doc, lines=3)

    # Metadata block
    meta_lines = [
        ("Date", date.today().strftime("%B %d, %Y")),
        ("Quality Score", f"{reflection.score}/10"),
        ("Status", "Approved" if reflection.passed else "Needs Review"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p.add_run(f"{label}: ")
        r_label.font.bold = True
        r_label.font.size = Pt(10)
        r_value = p.add_run(value)
        r_value.font.size = Pt(10)

    doc.add_page_break()


def _add_executive_summary(doc: Document, plan: ExecutionPlan) -> None:
    doc.add_heading("Executive Summary", level=1)

    p = doc.add_paragraph()
    p.add_run("Objective: ").bold = True
    p.add_run(plan.objective)

    if plan.assumptions:
        doc.add_heading("Key Assumptions", level=2)
        for assumption in plan.assumptions:
            doc.add_paragraph(assumption, style="List Bullet")

    doc.add_paragraph()


def _add_task_sections(doc: Document, plan: ExecutionPlan) -> None:
    # Skip the final "review & compilation" task — it consolidates; its content
    # is represented by the document itself.
    content_tasks = [t for t in plan.tasks if t.id < len(plan.tasks)]

    for task in content_tasks:
        doc.add_heading(f"{task.id}. {task.title}", level=1)
        _render_task_output(doc, task.output)
        doc.add_paragraph()

    # Add the final compilation task content as a standalone section if it has substance
    final_task = plan.tasks[-1]
    if final_task.output and len(final_task.output) > 50:
        doc.add_heading("Summary & Next Steps", level=1)
        _render_task_output(doc, final_task.output)


def _render_task_output(doc: Document, text: str) -> None:
    """
    Render task output text into the document.
    Detects basic markdown-like patterns:
      ## Heading -> Heading 2
      - bullet   -> List Bullet
      **bold**   -> bold run
      Plain text -> Normal paragraph
    """
    lines = text.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=2)
        elif line.startswith(("- ", "* ", "• ")):
            text_content = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, text_content)
        elif re.match(r"^\d+\.\s", line):
            text_content = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, text_content)
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, line)

        i += 1


def _add_inline_formatting(paragraph, text: str) -> None:
    """Handle **bold** and *italic* inline markdown."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|(.+?)(?=\*\*|\*|$))", re.DOTALL)
    pos = 0
    for match in re.finditer(r"(\*\*(.+?)\*\*|\*(.+?)\*)", text):
        # Text before this match
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        full = match.group(0)
        if full.startswith("**"):
            r = paragraph.add_run(match.group(2))
            r.bold = True
        else:
            r = paragraph.add_run(match.group(3))
            r.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_reflection_appendix(doc: Document, reflection: ReflectionResult) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix: Agent Quality Review", level=1)

    p = doc.add_paragraph()
    r = p.add_run(f"Quality Score: {reflection.score}/10  |  Status: ")
    r.bold = True
    status_run = p.add_run("PASSED" if reflection.passed else "NEEDS REVIEW")
    status_run.bold = True
    status_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) if reflection.passed else RGBColor(0xE7, 0x4C, 0x3C)

    if reflection.gaps:
        doc.add_heading("Identified Gaps", level=2)
        for gap in reflection.gaps:
            doc.add_paragraph(gap, style="List Bullet")

    if reflection.improvements:
        doc.add_heading("Suggested Improvements", level=2)
        for imp in reflection.improvements:
            doc.add_paragraph(imp, style="List Bullet")

    if not reflection.gaps and not reflection.improvements:
        doc.add_paragraph("No significant gaps or improvements identified. Document meets quality standards.")


def _add_footer(doc: Document, plan: ExecutionPlan) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run(f"{plan.document_title}  |  Generated by Autonomous AI Agent  |  {date.today().strftime('%Y-%m-%d')}")
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_SECONDARY

    # Page number field
    _add_page_number(p)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run("  |  Page ")
    run.font.size = Pt(8)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r.element.append(fldChar1)
    r.element.append(instrText)
    r.element.append(fldChar2)
    r.font.size = Pt(8)


def _add_spacer(doc: Document, lines: int = 1) -> None:
    for _ in range(lines):
        doc.add_paragraph()


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = BRAND_LIGHT


def _safe_filename(title: str) -> str:
    """Convert a document title to a safe filesystem name."""
    safe = re.sub(r"[^\w\s-]", "", title)
    safe = re.sub(r"\s+", "_", safe.strip())
    return safe[:80]
